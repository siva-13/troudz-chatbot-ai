"""
RAG Application Setup Script - Local Version
Helps set up and test the local RAG application without AWS
"""

import os
import subprocess
import sys

def check_environment():
    """Check if environment is properly set up"""
    print("🔍 Checking Environment Setup")
    print("=" * 30)

    # Check OpenAI API key
    openai_key = os.getenv('OPENAI_API_KEY')
    if openai_key and openai_key != 'your-openai-api-key-here':
        print("✅ OPENAI_API_KEY is set")
        return True
    else:
        print("❌ OPENAI_API_KEY is not set")
        print("\n🔧 To fix this:")
        print("Windows: set OPENAI_API_KEY=your-actual-openai-key")
        print("Mac/Linux: export OPENAI_API_KEY=your-actual-openai-key")
        return False

def test_imports():
    """Test if all required packages can be imported"""
    print("\n🧪 Testing Package Imports")
    print("=" * 30)

    packages = [
        ('fastapi', 'FastAPI web framework'),
        ('openai', 'OpenAI API client'),
        ('faiss', 'Vector similarity search'),
        ('tiktoken', 'Token counting'),
        ('docx', 'Word document processing'),
        ('numpy', 'Numerical computing'),
        ('uvicorn', 'ASGI server')
    ]

    all_good = True
    for package, description in packages:
        try:
            if package == 'docx':
                from docx import Document
            else:
                __import__(package)
            print(f"✅ {package:<10} - {description}")
        except ImportError as e:
            print(f"❌ {package:<10} - {description} (MISSING: {e})")
            all_good = False

    return all_good

def create_directories():
    """Create necessary directories"""
    print("\n📁 Creating Directories")
    print("=" * 25)

    dirs = ['./local_data', './data']

    for directory in dirs:
        try:
            os.makedirs(directory, exist_ok=True)
            print(f"✅ Created/verified: {directory}")
        except Exception as e:
            print(f"❌ Failed to create {directory}: {e}")

def test_local_server():
    """Test if the local server can be imported"""
    print("\n🖥️  Testing Server Import")
    print("=" * 25)

    try:
        # Try importing the local server
        import fastapi_server_local
        print("✅ Local server can be imported")
        return True
    except Exception as e:
        print(f"❌ Server import failed: {e}")
        return False

def create_sample_document():
    """Create a sample document for testing"""
    print("\n📄 Creating Sample Document")
    print("=" * 30)

    try:
        from docx import Document

        # Create sample document
        doc = Document()
        doc.add_heading('Sample RAG Document', 0)

        doc.add_paragraph('This is a sample document for testing the RAG application.')
        doc.add_paragraph('The RAG system can process Word documents like this one, extract text, create embeddings, and use the content to answer questions.')

        doc.add_heading('Key Features', level=1)
        doc.add_paragraph('1. Document ingestion and text extraction')
        doc.add_paragraph('2. Intelligent text chunking for optimal retrieval')
        doc.add_paragraph('3. OpenAI embeddings for semantic search')
        doc.add_paragraph('4. Conversation management with session tracking')
        doc.add_paragraph('5. Local storage for development and testing')

        doc.add_heading('Usage', level=1)
        doc.add_paragraph('Upload this document through the API and then ask questions about it. The system will find relevant sections and provide accurate answers based on the document content.')

        # Save the document
        sample_path = './sample_document.docx'
        doc.save(sample_path)
        print(f"✅ Created sample document: {sample_path}")
        return True

    except Exception as e:
        print(f"❌ Failed to create sample document: {e}")
        return False

def main():
    """Main setup function"""
    print("🚀 RAG Application Setup - Local Version")
    print("=" * 45)
    print("This setup script prepares your environment for the RAG application")
    print("LOCAL VERSION: No AWS credentials required!")
    print("=" * 45)

    # Check environment
    env_ok = check_environment()

    # Test imports
    imports_ok = test_imports()

    # Create directories
    create_directories()

    # Test server
    server_ok = test_local_server()

    # Create sample document
    doc_ok = create_sample_document()

    # Summary
    print("\n📊 Setup Summary")
    print("=" * 20)

    if env_ok and imports_ok and server_ok:
        print("🎉 Setup Complete! Your RAG application is ready.")
        print("\n🚀 To start the application:")
        print("   python fastapi_server_local.py")
        print("\n🌐 Then visit:")
        print("   http://localhost:8000/docs")
        print("\n📝 API Testing:")
        print("   1. Create a user via POST /users")
        print("   2. Upload the sample document via POST /documents/ingest")
        print("   3. Start a conversation via POST /conversations/start")
        print("   4. Send messages via POST /conversations/message")

        if doc_ok:
            print("\n📄 Sample document created: sample_document.docx")
            print("   Upload this through the API to test document ingestion")

    else:
        print("❌ Setup incomplete. Please fix the issues above.")

        if not env_ok:
            print("\n🔧 Set your OpenAI API key first:")
            print("   set OPENAI_API_KEY=your-actual-key")

        if not imports_ok:
            print("\n🔧 Install missing packages:")
            print("   pip install -r requirements_final_fix.txt")

    print("\n💡 Tips:")
    print("   - The local version stores data in ./local_data/")
    print("   - Vector indices are stored in ./data/")
    print("   - Check logs for any issues")
    print("   - Visit /health endpoint to check system status")

if __name__ == "__main__":
    main()
