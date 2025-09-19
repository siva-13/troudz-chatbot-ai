"""
Document Processing Module for RAG Application - FIXED VERSION
Handles Word document ingestion, text extraction, chunking, and embedding generation
"""

import os
import logging
from typing import List, Dict, Any, Tuple
from docx import Document
import tiktoken
import openai
import numpy as np
from datetime import datetime
import uuid
import json

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, openai_api_key: str, embedding_model: str = "text-embedding-3-small"):
        """
        Initialize the document processor

        Args:
            openai_api_key: OpenAI API key for embedding generation
            embedding_model: OpenAI embedding model to use
        """
        # Initialize OpenAI client with error handling for Pydantic issues
        try:
            self.openai_client = openai.OpenAI(api_key=openai_api_key)
        except Exception as e:
            logger.warning(f"OpenAI client initialization warning: {e}")
            # Try alternative initialization
            try:
                import openai as openai_module
                openai_module.api_key = openai_api_key
                self.openai_client = openai.OpenAI(api_key=openai_api_key)
            except Exception as e2:
                logger.error(f"Failed to initialize OpenAI client: {e2}")
                raise

        self.embedding_model = embedding_model
        self.tokenizer = tiktoken.get_encoding("cl100k_base")  # For GPT-3.5/4 models

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract raw text from Word document (.docx)
        Ignores formatting, headers, and footers

        Args:
            file_path: Path to the .docx file

        Returns:
            Raw text content as string
        """
        try:
            doc = Document(file_path)
            full_text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    full_text.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())

            return "\n".join(full_text)

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise

    def count_tokens(self, text: str) -> int:
        """
        Count tokens in text using tiktoken

        Args:
            text: Input text to count tokens for

        Returns:
            Number of tokens
        """
        return len(self.tokenizer.encode(text))

    def chunk_text(self, text: str, min_tokens: int = 500, max_tokens: int = 1000, overlap: int = 50) -> List[Dict[str, Any]]:
        """
        Chunk text into segments of 500-1000 tokens with overlap

        Args:
            text: Input text to chunk
            min_tokens: Minimum tokens per chunk
            max_tokens: Maximum tokens per chunk  
            overlap: Number of overlapping tokens between chunks

        Returns:
            List of chunks with metadata
        """
        sentences = text.split('. ')
        chunks = []
        current_chunk = []
        current_tokens = 0
        chunk_id = 0

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            sentence_tokens = self.count_tokens(sentence)

            # If adding this sentence would exceed max_tokens, create a new chunk
            if current_tokens + sentence_tokens > max_tokens and current_chunk:
                # Create chunk with metadata
                chunk_text = '. '.join(current_chunk)
                chunk_id += 1

                chunks.append({
                    'chunk_id': f"chunk_{chunk_id}",
                    'text': chunk_text,
                    'token_count': current_tokens,
                    'start_pos': len(chunks) * (max_tokens - overlap) if chunks else 0,
                    'end_pos': len(chunks) * (max_tokens - overlap) + current_tokens if chunks else current_tokens
                })

                # Start new chunk with overlap
                if overlap > 0 and len(current_chunk) > 1:
                    # Keep last few sentences for overlap
                    overlap_sentences = current_chunk[-2:]  # Keep last 2 sentences
                    current_chunk = overlap_sentences + [sentence]
                    current_tokens = sum(self.count_tokens(s) for s in overlap_sentences) + sentence_tokens
                else:
                    current_chunk = [sentence]
                    current_tokens = sentence_tokens
            else:
                current_chunk.append(sentence)
                current_tokens += sentence_tokens

        # Add final chunk if it meets minimum requirements
        if current_chunk and current_tokens >= min_tokens:
            chunk_text = '. '.join(current_chunk)
            chunk_id += 1

            chunks.append({
                'chunk_id': f"chunk_{chunk_id}",
                'text': chunk_text,
                'token_count': current_tokens,
                'start_pos': len(chunks) * (max_tokens - overlap) if chunks else 0,
                'end_pos': len(chunks) * (max_tokens - overlap) + current_tokens if chunks else current_tokens
            })

        logger.info(f"Created {len(chunks)} chunks from text")
        return chunks

    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """
        Generate embeddings using OpenAI Embeddings API with enhanced error handling

        Args:
            texts: List of text strings to embed

        Returns:
            List of embedding vectors
        """
        try:
            # Clean texts to avoid potential issues
            cleaned_texts = []
            for text in texts:
                if isinstance(text, str) and text.strip():
                    # Remove any problematic characters and limit length
                    cleaned_text = text.strip()[:8192]  # OpenAI max length
                    cleaned_texts.append(cleaned_text)
                else:
                    cleaned_texts.append("Empty text")

            # Create embeddings with better error handling
            response = self.openai_client.embeddings.create(
                model=self.embedding_model,
                input=cleaned_texts,
                encoding_format="float"  # Explicitly specify format
            )

            embeddings = [data.embedding for data in response.data]
            logger.info(f"Generated embeddings for {len(embeddings)} texts")
            return embeddings

        except Exception as e:
            logger.error(f"Error generating embeddings: {str(e)}")

            # Try alternative approach with individual requests
            logger.info("Trying individual embedding requests...")
            embeddings = []
            for i, text in enumerate(texts):
                try:
                    response = self.openai_client.embeddings.create(
                        model=self.embedding_model,
                        input=[text[:8192]],  # Limit text length
                        encoding_format="float"
                    )
                    embeddings.append(response.data[0].embedding)
                    logger.info(f"Generated embedding {i+1}/{len(texts)}")
                except Exception as e2:
                    logger.error(f"Error generating embedding {i+1}: {e2}")
                    # Use zero vector as fallback
                    embeddings.append([0.0] * 1536)  # Default dimension for text-embedding-3-small

            if embeddings:
                logger.info(f"Generated {len(embeddings)} embeddings (some may be fallback)")
                return embeddings
            else:
                raise Exception(f"Failed to generate embeddings: {e}")

    def process_document(self, file_path: str, document_id: str = None) -> Dict[str, Any]:
        """
        Complete document processing pipeline with enhanced error handling

        Args:
            file_path: Path to the Word document
            document_id: Optional document ID (generated if not provided)

        Returns:
            Dictionary containing processed document data
        """
        if document_id is None:
            document_id = str(uuid.uuid4())

        try:
            # Extract text
            logger.info(f"Processing document: {file_path}")
            raw_text = self.extract_text_from_docx(file_path)

            if not raw_text.strip():
                raise Exception("Document appears to be empty or contains no readable text")

            # Chunk text
            chunks = self.chunk_text(raw_text)

            if not chunks:
                raise Exception("Failed to create chunks from document text")

            # Generate embeddings for all chunks
            chunk_texts = [chunk['text'] for chunk in chunks]
            embeddings = self.generate_embeddings(chunk_texts)

            if len(embeddings) != len(chunks):
                logger.warning(f"Embedding count ({len(embeddings)}) doesn't match chunk count ({len(chunks)})")

            # Combine chunks with embeddings
            for i, chunk in enumerate(chunks):
                if i < len(embeddings):
                    chunk['embedding'] = embeddings[i]
                    chunk['document_id'] = document_id
                else:
                    # Fallback embedding if needed
                    chunk['embedding'] = [0.0] * 1536
                    chunk['document_id'] = document_id
                    logger.warning(f"Used fallback embedding for chunk {i+1}")

            result = {
                'document_id': document_id,
                'filename': os.path.basename(file_path),
                'raw_text': raw_text,
                'total_tokens': self.count_tokens(raw_text),
                'chunks': chunks,
                'processed_at': datetime.now().isoformat()
            }

            logger.info(f"Successfully processed document: {document_id} with {len(chunks)} chunks")
            return result

        except Exception as e:
            logger.error(f"Error processing document: {e}")
            raise
